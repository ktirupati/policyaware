from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "PolicyAware_AI_Gateway_User_Guide.docx"


BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F3F6F8"
DARK = "1F2933"
GREEN = "DDEFE5"
AMBER = "FFF2CC"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_code(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    for line_no, line in enumerate(code.strip("\n").splitlines()):
        if line_no:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(8.5)


def add_note(doc: Document, title: str, body: str, fill: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p.add_run("\n" + body)


def add_kv_table(doc: Document, title: str, rows: list[tuple[str, str]]) -> None:
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Item", True, "FFFFFF")
    set_cell_text(hdr[1], "Details", True, "FFFFFF")
    shade_cell(hdr[0], BLUE)
    shade_cell(hdr[1], BLUE)
    set_repeat_table_header(table.rows[0])
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, True)
        set_cell_text(cells[1], value)


def add_command_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Command", True, "FFFFFF")
    set_cell_text(hdr[1], "Purpose", True, "FFFFFF")
    shade_cell(hdr[0], BLUE)
    shade_cell(hdr[1], BLUE)
    set_repeat_table_header(table.rows[0])
    for command, purpose in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], command)
        for p in cells[0].paragraphs:
            for run in p.runs:
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                run.font.size = Pt(8)
        set_cell_text(cells[1], purpose)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.08

    for name, size in [("Title", 26), ("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def build() -> None:
    doc = Document()
    setup_styles(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    doc.core_properties.title = "PolicyAware AI Gateway User Guide"
    doc.core_properties.subject = "Installation, usage, and working examples"
    doc.core_properties.author = "PolicyAware AI Gateway Contributors"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PolicyAware AI Gateway")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("User Guide, Installation Steps, and Working Examples")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(DARK)

    add_note(
        doc,
        "What this framework does",
        "PolicyAware is a governed AI execution layer. YAML files define the rules, but the framework does "
        "not only check YAML. It checks live user prompts, request context, tool calls, model responses, "
        "routing decisions, and audit evidence for LLM, RAG, agent, and MCP-style tool workflows.",
        GREEN,
    )

    add_kv_table(
        doc,
        "Package Identity",
        [
            ("Framework name", "PolicyAware AI Gateway"),
            ("Python package", "policyaware"),
            ("Python import", "import policyaware"),
            ("CLI command", "policyaware"),
            ("Current version", "0.2.0"),
            ("License", "Apache-2.0"),
            ("Provider adapters", "Simulated provider and OpenAI-compatible HTTP adapter"),
            ("Audit storage", "JSONL by default, SQLite for persistent local storage"),
            ("Observability", "Prometheus text export and OpenTelemetry-shaped JSON export"),
        ],
    )

    doc.add_page_break()

    doc.add_heading("1. Quick Installation", level=1)
    doc.add_paragraph("Use the local editable install while developing from this repository:")
    add_code(doc, 'pip install -e ".[dev]"')
    doc.add_paragraph("After the package is published to PyPI under the policyaware name, users can install it with:")
    add_code(doc, "pip install policyaware")
    add_note(
        doc,
        "Important",
        "The public pip install command works only after the package is uploaded to PyPI. Until then, "
        "run the editable install from the repository root.",
        AMBER,
    )

    doc.add_heading("2. What PolicyAware Checks", level=1)
    add_note(
        doc,
        "YAML is configuration, not the only thing being checked",
        "Policy files such as examples/policies/basic.yaml define the rules. At runtime, PolicyAware applies "
        "those rules to real AI requests, prompts, context, tool calls, model outputs, and audit traces.",
        AMBER,
    )
    checks = doc.add_table(rows=1, cols=4)
    checks.style = "Table Grid"
    checks.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = checks.rows[0].cells
    for cell, text in zip(hdr, ["Area Checked", "What It Looks At", "Examples", "Possible Outcome"]):
        set_cell_text(cell, text, True, "FFFFFF")
        shade_cell(cell, BLUE)
    set_repeat_table_header(checks.rows[0])
    for row in [
        (
            "User prompts / messages",
            "Message text sent by users or agents before it reaches a model.",
            "PII, PHI, secrets, API keys, emails, phone numbers, sensitive business text.",
            "Deny, redact, classify as higher risk, or allow.",
        ),
        (
            "Request context",
            "Metadata supplied by the application around the request.",
            "User role, tenant, region, task type, risk level, domain, autonomy level.",
            "Apply RBAC, tenant isolation, region restrictions, or approval requirements.",
        ),
        (
            "YAML policies",
            "Rules that define what the organization allows or blocks.",
            "basic.yaml, regulated-rag.yaml, tool-governance.yaml.",
            "Produce allow, deny, conditional_allow, or require_approval decisions.",
        ),
        (
            "Tool calls",
            "Agent or MCP-style connector actions before execution.",
            "Connector name, action name, user role, arguments, write/delete side effects.",
            "Allow read actions, deny destructive actions, require approval for writes.",
        ),
        (
            "Model responses",
            "Text returned by the model before it is trusted by the application.",
            "Sensitive data leakage, missing citations, policy consistency.",
            "Flag eval failure, record safety score, or block in future stricter flows.",
        ),
    ]:
        cells = checks.add_row().cells
        for idx, text in enumerate(row):
            set_cell_text(cells[idx], text, idx == 0)

    doc.add_heading("3. Runtime Decision Flow", level=1)
    flow = doc.add_table(rows=1, cols=3)
    flow.style = "Table Grid"
    flow.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = flow.rows[0].cells
    for cell, text in zip(hdr, ["Step", "Runtime Check", "Example"]):
        set_cell_text(cell, text, True, "FFFFFF")
        shade_cell(cell, BLUE)
    set_repeat_table_header(flow.rows[0])
    for row in [
        ("1", "Inspect prompt/message text", "Detect jane@example.com as PII."),
        ("2", "Classify risk", "Healthcare + PHI + agentic autonomy becomes high risk."),
        ("3", "Evaluate YAML policy", "block_secrets denies secret-like prompt content."),
        ("4", "Apply transform if allowed", "Redact PII before sending to model."),
        ("5", "Route to model", "High-risk requests prefer higher-quality compliant models."),
        ("6", "Evaluate model response", "Check sensitive leakage and citation requirements."),
        ("7", "Write audit trace", "Save decision, reason codes, request snapshot, and eval scores."),
    ]:
        cells = flow.add_row().cells
        for idx, text in enumerate(row):
            set_cell_text(cells[idx], text, idx == 0)

    doc.add_heading("4. Working Examples For Each Checked Area", level=1)
    doc.add_paragraph(
        "The table above summarizes what PolicyAware checks. The examples below show how each category works "
        "with real CLI commands or Python code."
    )

    doc.add_heading("4.1 User Prompts And Messages", level=2)
    doc.add_paragraph(
        "PolicyAware inspects the actual user or agent message before it reaches a model. This is where PII, PHI, "
        "secrets, and other sensitive content are detected."
    )
    add_code(
        doc,
        '''
policyaware policy explain examples/policies/basic.yaml ^
  --role support_agent ^
  --prompt "Email jane@example.com about claim ACME-42"
''',
    )
    add_note(
        doc,
        "Expected outcome",
        "The prompt contains an email address. PolicyAware detects PII, applies the YAML policy, and returns "
        "conditional_allow with a redact action for non-privileged users.",
    )
    add_code(
        doc,
        '''
{
  "decision": "conditional_allow",
  "reason_codes": [
    "DATA.PII_DETECTED",
    "POLICY.ALLOW_MATCHED",
    "POLICY.TRANSFORM_APPLIED"
  ],
  "remediation": ["Transforms were applied before execution."]
}
''',
    )
    doc.add_paragraph("Secret detection example:")
    add_code(
        doc,
        '''
policyaware policy explain examples/policies/basic.yaml ^
  --role developer ^
  --prompt "Use secret_api_key_abcdefghijklmnop in the deployment"
''',
    )
    add_note(
        doc,
        "Expected outcome",
        "The request is denied before model execution because a secret-like value was found.",
        AMBER,
    )

    doc.add_heading("4.2 Request Context", level=2)
    doc.add_paragraph(
        "PolicyAware checks metadata supplied by the application, not only the prompt text. Context is how the "
        "gateway understands role, tenant, region, task type, domain, autonomy level, and business impact."
    )
    add_code(
        doc,
        '''
from policyaware import Gateway, GatewayRequest

gateway = Gateway.from_policy_file("examples/policies/basic.yaml")

response = gateway.chat(
    GatewayRequest(
        tenant="acme",
        app="claims-assistant",
        user={"id": "u_456", "role": "claims_adjuster"},
        context={
            "region": "us",
            "task_type": "summarization",
            "risk": "low",
            "domain": "insurance",
            "autonomy": "assistive",
            "business_impact": "medium",
        },
        messages=[{"role": "user", "content": "Summarize claim ACME-42"}],
    )
)

print(response.policy.decision)
print(response.policy.risk_tier)
print(response.policy.reason_codes)
''',
    )
    add_bullets(
        doc,
        [
            "Role check: claims_adjuster is allowed by the sample policy.",
            "Region check: us matches the policy requirement.",
            "Task check: summarization can be used for routing, evals, and audit.",
            "Risk check: low risk allows execution without approval.",
        ],
    )

    doc.add_heading("4.3 YAML Policies", level=2)
    doc.add_paragraph(
        "YAML files define the rules, but they are evaluated against live request data. They are not the only "
        "thing being checked; they are the policy source used by the runtime decision engine."
    )
    add_code(
        doc,
        '''
rules:
  - name: block_secrets
    effect: deny
    when:
      data.contains_secrets: true

  - name: redact_pii_for_non_privileged_users
    effect: transform
    action: redact
    when:
      data.contains_pii: true
      user.role_not_in: ["privacy_admin", "compliance_officer"]
''',
    )
    doc.add_paragraph("These rules are applied to runtime facts such as:")
    add_code(
        doc,
        '''
{
  "data": {
    "contains_pii": true,
    "contains_secrets": false
  },
  "user": {
    "role": "support_agent"
  },
  "request": {
    "region": "us",
    "task_type": "support"
  }
}
''',
    )
    add_note(
        doc,
        "Expected outcome",
        "The YAML policy does not act alone. It uses runtime facts from the prompt, user, tenant, request context, "
        "risk classifier, and data protection engine.",
    )

    doc.add_heading("4.4 Tool Calls", level=2)
    doc.add_paragraph(
        "For agents and MCP-style connectors, PolicyAware checks connector name, action name, user role, arguments, "
        "side effects, and whether approval is required."
    )
    add_code(
        doc,
        '''
policyaware tools check examples/policies/tool-governance.yaml ^
  --agent code_assistant ^
  --connector github ^
  --action read_file ^
  --role developer
''',
    )
    doc.add_paragraph("Expected outcome for a read-only action:")
    add_code(
        doc,
        '''
{
  "decision": "allow",
  "connector_id": "github",
  "action": "read_file",
  "reason_codes": ["TOOL.RATE_LIMIT_DECLARED", "TOOL.ALLOWED"]
}
''',
    )
    doc.add_paragraph("Expected outcome for a write action:")
    add_code(
        doc,
        '''
policyaware tools check examples/policies/tool-governance.yaml ^
  --agent code_assistant ^
  --connector github ^
  --action create_pr ^
  --role developer
''',
    )
    add_code(
        doc,
        '''
{
  "decision": "require_approval",
  "connector_id": "github",
  "action": "create_pr",
  "approval_required": true,
  "reason_codes": ["TOOL.APPROVAL_REQUIRED"]
}
''',
    )
    add_note(
        doc,
        "Why this matters",
        "Tool calls are where an AI system can take action. PolicyAware treats tool actions as governed operations, "
        "especially when they can write, delete, deploy, query sensitive systems, or communicate externally.",
        AMBER,
    )

    doc.add_heading("4.5 Model Responses", level=2)
    doc.add_paragraph(
        "PolicyAware also evaluates the response after the model returns. This catches sensitive data leakage, "
        "missing citations, and responses that are inconsistent with policy decisions."
    )
    add_code(
        doc,
        '''
response = gateway.chat(
    GatewayRequest(
        tenant="acme",
        app="rag-assistant",
        user={"id": "u_789", "role": "analyst"},
        context={
            "region": "us",
            "task_type": "rag_answer",
            "risk": "medium",
            "require_citations": True,
        },
        messages=[{"role": "user", "content": "Answer using the retrieved policy docs."}],
    )
)

for result in response.evals:
    print(result.name, result.passed, result.score, result.reason)
''',
    )
    doc.add_paragraph("Possible evaluation results:")
    add_code(
        doc,
        '''
sensitive_data_leakage True 1.0 No sensitive data detected in output.
citation_required False 0.0 Citation marker missing.
policy_compliance True 1.0 Output is consistent with policy decision.
''',
    )
    add_note(
        doc,
        "Expected outcome",
        "If require_citations is true and the model output has no citation marker, the citation_required eval fails. "
        "The result is included in the response object and audit trace.",
    )

    doc.add_heading("5. Core Request Lifecycle", level=1)
    add_bullets(
        doc,
        [
            "The application sends a GatewayRequest to PolicyAware.",
            "The data protection engine checks for PII, PHI, secrets, and sensitive content.",
            "The risk classifier maps the request to low, medium, high, or critical risk.",
            "The policy engine makes an allow, deny, conditional_allow, or require_approval decision.",
            "Allowed requests are routed to an approved model or simulated provider.",
            "Runtime evaluation checks leakage, citation requirements, and policy consistency.",
            "Audit traces are written for replay, incident review, and compliance evidence.",
        ],
    )

    doc.add_heading("6. Smoke Test Commands", level=1)
    add_command_table(
        doc,
        [
            ("policyaware dev simulate", "Run built-in local governance scenarios."),
            (
                'policyaware policy explain examples/policies/basic.yaml --prompt "Email jane@example.com"',
                "Explain a PII redaction decision.",
            ),
            (
                'policyaware risk classify "Review patient id ABCDE diagnosis: flu" --domain healthcare --autonomy agentic',
                "Classify a healthcare-style request as higher risk.",
            ),
            (
                "policyaware tools check examples/policies/tool-governance.yaml --agent code_assistant --connector github --action create_pr",
                "Check whether a governed tool action requires approval.",
            ),
            ("policyaware eval run examples/evals/governance_cases.yaml", "Parse the governance eval suite."),
        ],
    )

    doc.add_heading("7. Python SDK Example", level=1)
    doc.add_paragraph("This example sends a support-copilot request through the gateway.")
    add_code(
        doc,
        '''
from policyaware import Gateway, GatewayRequest

gateway = Gateway.from_policy_file("examples/policies/basic.yaml")

response = gateway.chat(
    GatewayRequest(
        tenant="acme",
        app="support-copilot",
        user={"id": "u_123", "role": "support_agent"},
        context={
            "region": "us",
            "task_type": "support",
            "risk": "low",
            "domain": "support",
        },
        messages=[
            {"role": "user", "content": "Email jane@example.com about the claim."}
        ],
    )
)

print(response.content)
print(response.policy.decision)
print(response.policy.risk_tier)
print(response.policy.reason_codes)
print(response.policy.explanation.summary)
print(response.trace_id)
''',
    )
    add_note(
        doc,
        "Expected result",
        "The email address is detected as PII. The support role is allowed by policy, the request is "
        "conditionally allowed with redaction, the simulated provider receives redacted text, and an audit trace is written.",
    )

    doc.add_heading("8. Runtime Prompt And Context Example", level=1)
    doc.add_paragraph(
        "This example shows that the framework checks both prompt content and request context. "
        "The YAML policy is only the rulebook; the live GatewayRequest is what gets evaluated."
    )
    add_code(
        doc,
        '''
request = GatewayRequest(
    tenant="acme",
    app="claims-assistant",
    user={"id": "u_456", "role": "claims_adjuster"},
    context={
        "region": "us",
        "task_type": "summarization",
        "domain": "insurance",
        "autonomy": "assistive",
        "business_impact": "medium",
    },
    messages=[
        {"role": "user", "content": "Summarize claim for jane@example.com"}
    ],
)
''',
    )
    add_bullets(
        doc,
        [
            "Prompt check: detects jane@example.com as PII.",
            "Context check: reads tenant acme, role claims_adjuster, region us, task type summarization, and domain insurance.",
            "Policy check: matches allow/redact rules from YAML.",
            "Outcome: conditional_allow with redact action and an audit trace.",
        ],
    )

    doc.add_heading("9. Policy Testing And Explanation", level=1)
    add_code(
        doc,
        '''
policyaware policy test examples/policies/basic.yaml ^
  --role support_agent ^
  --risk low ^
  --prompt "Summarize this customer request."
''',
    )
    doc.add_paragraph("Use policy explain when you need machine-readable reason codes and remediation guidance:")
    add_code(
        doc,
        '''
policyaware policy explain examples/policies/basic.yaml ^
  --role support_agent ^
  --prompt "Email jane@example.com about the claim."
''',
    )
    doc.add_paragraph("Expected explanation shape:")
    add_code(
        doc,
        '''
{
  "decision": "conditional_allow",
  "summary": "conditional_allow: Allowed with transforms.",
  "reason_codes": [
    "DATA.PII_DETECTED",
    "RISK.MEDIUM",
    "POLICY.ALLOW_MATCHED",
    "POLICY.TRANSFORM_APPLIED"
  ],
  "matched_policy_ids": [
    "allow_low_medium_risk_enterprise_users",
    "redact_pii_for_non_privileged_users"
  ],
  "remediation": ["Transforms were applied before execution."]
}
''',
    )

    doc.add_heading("10. Deny Secrets Before Model Execution", level=1)
    add_code(
        doc,
        '''
policyaware policy explain examples/policies/basic.yaml ^
  --role developer ^
  --prompt "Use secret_api_key_abcdefghijklmnop in the deployment."
''',
    )
    add_note(
        doc,
        "Expected result",
        "Secret-like content is detected, the block_secrets policy rule matches, and the request is denied before model execution. "
        "Reason codes include DATA.SECRET_DETECTED and POLICY.DENY_MATCHED.",
        AMBER,
    )

    doc.add_heading("11. Risk Classification", level=1)
    add_code(
        doc,
        '''
policyaware risk classify "Review patient id ABCDE diagnosis: flu" ^
  --role analyst ^
  --domain healthcare ^
  --autonomy agentic ^
  --action-type read
''',
    )
    add_kv_table(
        doc,
        "Risk Tiers",
        [
            ("low", "Routine, low-sensitivity requests. Basic logging and evaluation are enough."),
            ("medium", "Internal or customer-support requests with moderate sensitivity. Redaction may be required."),
            ("high", "Regulated, sensitive, or tool-using workflows. Strict evaluation and approval hooks may apply."),
            ("critical", "Autonomous or high-impact actions. Default fail-safe is denial or mandatory approval."),
        ],
    )

    doc.add_heading("12. MCP-Style Tool Governance", level=1)
    doc.add_paragraph(
        "Tool governance controls which agents can call which connectors and actions. This is useful for MCP-style "
        "tool servers, internal connectors, code assistants, analytics agents, and multi-agent workflows."
    )
    add_code(
        doc,
        '''
policyaware tools check examples/policies/tool-governance.yaml ^
  --agent code_assistant ^
  --connector github ^
  --action create_pr ^
  --role developer
''',
    )
    add_note(
        doc,
        "Expected result",
        "Creating a pull request is a write action, so the sample policy returns require_approval.",
    )
    add_code(
        doc,
        '''
{
  "decision": "require_approval",
  "connector_id": "github",
  "action": "create_pr",
  "approval_required": true,
  "reason_codes": ["TOOL.APPROVAL_REQUIRED"]
}
''',
    )

    doc.add_heading("13. Tool Call Runtime Data", level=1)
    doc.add_paragraph(
        "Tool governance checks the requested connector, action, arguments, user role, and approval requirements before execution."
    )
    add_code(
        doc,
        '''
ToolCallRequest(
    agent_id="code_assistant",
    connector_id="github",
    action="create_pr",
    user={"id": "u_123", "role": "developer"},
    arguments={"repository": "acme/app", "branch": "policy-fix"},
)
''',
    )
    add_bullets(
        doc,
        [
            "Connector check: github is known in tool-governance.yaml.",
            "Action check: create_pr is configured as require_approval.",
            "Role check: developer is allowed to request the action.",
            "Approval check: write action returns require_approval before execution.",
        ],
    )

    doc.add_heading("14. Model Response Evaluation", level=1)
    doc.add_paragraph(
        "PolicyAware also checks model responses after execution. The MVP includes leakage, citation, and policy-consistency checks."
    )
    response_checks = doc.add_table(rows=1, cols=3)
    response_checks.style = "Table Grid"
    hdr = response_checks.rows[0].cells
    for cell, text in zip(hdr, ["Response Check", "What It Detects", "Why It Matters"]):
        set_cell_text(cell, text, True, "FFFFFF")
        shade_cell(cell, BLUE)
    set_repeat_table_header(response_checks.rows[0])
    for row in [
        ("Sensitive data leakage", "PII, PHI, secrets in output", "Prevents unsafe data exposure."),
        ("Citation requirements", "Missing citation markers when require_citations is true", "Improves RAG trust and reviewability."),
        ("Policy consistency", "Output returned despite denied policy decision", "Catches governance path bugs."),
    ]:
        cells = response_checks.add_row().cells
        for idx, text in enumerate(row):
            set_cell_text(cells[idx], text, idx == 0)

    doc.add_heading("15. Evaluation Suite Example", level=1)
    add_code(doc, "policyaware eval run examples/evals/governance_cases.yaml")
    doc.add_paragraph(
        "The MVP eval runner parses governance eval definitions and returns a report-shaped object. "
        "The next production step is full model-backed execution against golden datasets."
    )
    add_code(
        doc,
        '''
{
  "suite": "governance_policy_eval",
  "checks": 4,
  "cases": 2,
  "status": "configured",
  "report": {
    "policy_compliance_score": 1.0,
    "safety_score": 1.0
  }
}
''',
    )

    doc.add_heading("16. Audit, Bundle, And Replay", level=1)
    doc.add_paragraph("Run a request to create an audit trace:")
    add_code(doc, 'policyaware chat examples/policies/basic.yaml "Summarize this customer ticket." --role support_agent --risk low')
    doc.add_paragraph("Traces are written to:")
    add_code(doc, ".policyaware/traces.jsonl")
    doc.add_paragraph("Generate an audit bundle from a trace ID:")
    add_code(
        doc,
        '''
policyaware audit bundle trc_your_trace_id ^
  --traces-file .policyaware/traces.jsonl ^
  --out .policyaware/audit-bundle
''',
    )
    doc.add_paragraph("Replay a trace against a policy:")
    add_code(
        doc,
        '''
policyaware audit replay trc_your_trace_id ^
  examples/policies/basic.yaml ^
  --traces-file .policyaware/traces.jsonl
''',
    )
    add_bullets(
        doc,
        [
            "trace.json contains the complete replayable audit record.",
            "decision.json contains the policy decision and reason codes.",
            "request.json contains the request snapshot.",
            "eval_report.json contains runtime evaluation results.",
            "summary.md gives reviewers a readable audit summary.",
        ],
    )

    doc.add_heading("17. Real Provider Adapters", level=1)
    doc.add_paragraph(
        "PolicyAware now includes a real provider adapter foundation. The local simulated provider remains the "
        "default for demos and tests, while OpenAICompatibleProvider can call hosted or self-hosted endpoints "
        "that implement the chat completions API shape."
    )
    add_code(
        doc,
        '''
from policyaware import Gateway, OpenAICompatibleProvider, ProviderRegistry

provider = OpenAICompatibleProvider(
    base_url="https://api.example.com/v1",
    api_key="YOUR_TOKEN",
)

registry = ProviderRegistry({
    "openai-compatible": provider
})

gateway = Gateway.from_policy_file("examples/policies/basic.yaml")
gateway.provider_registry = registry
''',
    )
    add_note(
        doc,
        "How routing connects to providers",
        "The model router returns a model candidate with a provider name. ProviderRegistry uses that provider name "
        "to select the adapter. If no matching provider is registered, the local simulated provider remains available.",
    )
    add_bullets(
        doc,
        [
            "Use SimulatedProvider for local development without API keys.",
            "Use OpenAICompatibleProvider for OpenAI-compatible hosted or local endpoints.",
            "Add future adapters by implementing the ModelProvider interface.",
        ],
    )

    doc.add_heading("18. Persistent Audit Storage And Trace Viewer", level=1)
    doc.add_paragraph(
        "PolicyAware supports JSONL audit traces and now also includes SQLiteAuditLogger for persistent local storage. "
        "A static HTML trace viewer can be generated from either JSONL or SQLite."
    )
    add_code(
        doc,
        '''
from policyaware import Gateway
from policyaware.audit import SQLiteAuditLogger

gateway = Gateway.from_policy_file("examples/policies/basic.yaml")
gateway.audit_logger = SQLiteAuditLogger(".policyaware/audit.db")
''',
    )
    doc.add_paragraph("Generate a static viewer from JSONL traces:")
    add_code(
        doc,
        '''
policyaware audit view ^
  --traces-file .policyaware/traces.jsonl ^
  --out .policyaware/trace-viewer.html
''',
    )
    doc.add_paragraph("Generate a static viewer from SQLite audit storage:")
    add_code(
        doc,
        '''
policyaware audit view-sqlite ^
  --db .policyaware/audit.db ^
  --out .policyaware/trace-viewer.html
''',
    )
    add_note(
        doc,
        "Trace viewer contents",
        "The viewer shows trace ID, tenant, app, decision, risk tier, model, latency, and reason codes. "
        "It is intentionally static so it can be opened locally or attached as a review artifact.",
    )

    doc.add_heading("19. OpenTelemetry And Prometheus Exporters", level=1)
    doc.add_paragraph(
        "The observability exporters convert local audit traces into formats that operations teams already understand."
    )
    add_command_table(
        doc,
        [
            (
                "policyaware observability prometheus --traces-file .policyaware/traces.jsonl --out .policyaware/metrics.prom",
                "Export Prometheus text exposition metrics.",
            ),
            (
                "policyaware observability otel-json --traces-file .policyaware/traces.jsonl --out .policyaware/otel-spans.json",
                "Export OpenTelemetry-shaped JSON span records.",
            ),
        ],
    )
    doc.add_paragraph("Example Prometheus output:")
    add_code(
        doc,
        '''
policyaware_requests_total 12
policyaware_policy_decisions_total{decision="allow"} 8
policyaware_policy_decisions_total{decision="deny"} 3
policyaware_risk_tiers_total{tier="high"} 2
policyaware_latency_ms_sum 1840
''',
    )
    doc.add_paragraph("Example OpenTelemetry-shaped span:")
    add_code(
        doc,
        '''
{
  "name": "policyaware.gateway.request",
  "trace_id": "trc_123",
  "attributes": {
    "policyaware.tenant": "acme",
    "policyaware.app": "support-copilot",
    "policyaware.decision": "conditional_allow",
    "policyaware.risk_tier": "medium",
    "policyaware.reason_codes": ["DATA.PII_DETECTED", "POLICY.TRANSFORM_APPLIED"]
  },
  "duration_ms": 42
}
''',
    )

    doc.add_heading("20. Approval Integrations", level=1)
    doc.add_paragraph(
        "When a policy returns require_approval, PolicyAware now creates an approval request instead of executing the model. "
        "The MVP includes a no-op local client, a file-backed queue, and a webhook integration."
    )
    add_code(
        doc,
        '''
from policyaware.approvals import FileApprovalClient

gateway.approval_client = FileApprovalClient(".policyaware/approvals.jsonl")
''',
    )
    add_code(
        doc,
        '''
from policyaware.approvals import WebhookApprovalClient

gateway.approval_client = WebhookApprovalClient(
    "https://workflow.example.com/ai-approval"
)
''',
    )
    add_note(
        doc,
        "Expected behavior",
        "For approval-required requests, the gateway returns an approval ID, records the audit trace, and does not call the model.",
        AMBER,
    )
    add_code(
        doc,
        '''
{
  "approval_id": "apr_123",
  "tenant": "acme",
  "app": "claims-assistant",
  "status": "pending",
  "decision": {
    "decision": "require_approval",
    "reason_codes": ["POLICY.APPROVAL_REQUIRED"]
  }
}
''',
    )

    doc.add_heading("21. Executable Golden Dataset Evaluations", level=1)
    doc.add_paragraph(
        "The eval runner can now execute golden cases through a real Gateway and compare actual policy decisions "
        "and reason codes against expected governance outcomes."
    )
    add_code(
        doc,
        '''
policyaware eval run examples/evals/executable_governance_cases.yaml ^
  --policy-file examples/policies/basic.yaml
''',
    )
    doc.add_paragraph("Example golden case:")
    add_code(
        doc,
        '''
cases:
  - id: secret_is_denied
    input: "Use secret_api_key_abcdefghijklmnop in the deployment."
    user:
      id: eval_user
      role: developer
    context:
      region: us
      risk: low
      task_type: code_assistant
    expected:
      decision: deny
      reason_codes:
        - DATA.SECRET_DETECTED
        - POLICY.DENY_MATCHED
''',
    )
    doc.add_paragraph("Example report shape:")
    add_code(
        doc,
        '''
{
  "suite": "executable_governance_policy_eval",
  "cases": 3,
  "status": "configured",
  "report": {
    "cases": 3,
    "passed": 3,
    "failed": 0,
    "policy_compliance_score": 1.0,
    "safety_score": 1.0
  }
}
''',
    )

    doc.add_heading("22. Updated Production Commands", level=1)
    add_command_table(
        doc,
        [
            ("policyaware audit view", "Create a static trace viewer from JSONL traces."),
            ("policyaware audit view-sqlite", "Create a static trace viewer from SQLite traces."),
            ("policyaware observability prometheus", "Export Prometheus metrics."),
            ("policyaware observability otel-json", "Export OpenTelemetry-shaped JSON spans."),
            (
                "policyaware eval run examples/evals/executable_governance_cases.yaml --policy-file examples/policies/basic.yaml",
                "Run executable governance golden dataset checks.",
            ),
        ],
    )

    doc.add_heading("23. Example Enterprise Policy", level=1)
    add_code(
        doc,
        '''
id: basic_enterprise_policy
default: deny

rules:
  - name: critical_requires_approval
    effect: require_approval
    when:
      risk.tier: "critical"

  - name: block_secrets
    effect: deny
    when:
      data.contains_secrets: true

  - name: allow_low_medium_risk_enterprise_users
    effect: allow
    when:
      user.role_in: ["support_agent", "claims_adjuster", "developer", "privacy_admin"]
      request.risk_in: ["low", "medium"]
      request.region: "us"

  - name: redact_pii_for_non_privileged_users
    effect: transform
    action: redact
    when:
      data.contains_pii: true
      user.role_not_in: ["privacy_admin", "compliance_officer"]

  - name: require_approval_for_high_risk
    effect: require_approval
    when:
      risk.tier_in: ["high", "critical"]
''',
    )

    doc.add_heading("24. Example Tool Governance Policy", level=1)
    add_code(
        doc,
        '''
id: mcp_tool_governance
schema_version: "0.2"
default: deny

connectors:
  - id: github
    type: mcp
    risk: medium
    actions:
      read_file:
        effect: allow
        risk: low
        side_effect: none
        when:
          user.role_in: ["developer", "security_engineer"]
        limits:
          calls_per_minute: 60

      create_pr:
        effect: require_approval
        risk: high
        side_effect: write
        when:
          user.role_in: ["developer", "maintainer"]

      delete_branch:
        effect: deny
        risk: critical
        side_effect: delete
''',
    )

    doc.add_heading("25. Real-Time Use Cases", level=1)
    use_cases = doc.add_table(rows=1, cols=3)
    use_cases.style = "Table Grid"
    hdr = use_cases.rows[0].cells
    for cell, text in zip(hdr, ["Scenario", "What PolicyAware Enforces", "Example Outcome"]):
        set_cell_text(cell, text, True, "FFFFFF")
        shade_cell(cell, BLUE)
    set_repeat_table_header(use_cases.rows[0])
    for row in [
        ("Customer support copilot", "PII redaction, role-based access, audit logging", "Conditional allow with redaction"),
        ("Regulated RAG assistant", "PHI detection, citation requirement, higher risk tier", "Approval or strict route"),
        ("Code assistant", "Secret blocking and GitHub action approval", "Deny secret prompt; approve PR creation"),
        ("Analytics agent", "Database/action restrictions and row limits", "Deny restricted database query"),
        ("Multi-agent workflow", "Tool side-effect checks and approval hooks", "Planner allowed, executor gated"),
    ]:
        cells = use_cases.add_row().cells
        for idx, text in enumerate(row):
            set_cell_text(cells[idx], text, idx == 0)

    doc.add_heading("26. Current MVP Limitations", level=1)
    add_bullets(
        doc,
        [
            "Provider-specific adapters beyond OpenAI-compatible APIs are still future work.",
            "MCP support currently checks policy but is not a full MCP proxy.",
            "Rate and budget limits are declared but not yet backed by persistent enforcement.",
            "The trace viewer is static HTML rather than a hosted dashboard.",
            "OpenTelemetry support is currently an OpenTelemetry-shaped JSON export, not native SDK emission.",
            "Golden evals execute policy outcomes, but model-backed RAG grounding and hallucination judges are still future work.",
            "Approval workflows support file and webhook hooks, but packaged Slack, Jira, ServiceNow, and email connectors are future work.",
        ],
    )

    doc.add_heading("27. Recommended Next Steps", level=1)
    add_bullets(
        doc,
        [
            "Publish the package to PyPI after confirming the policyaware name is available.",
            "Add provider-specific adapters for your target model platforms.",
            "Replace the static trace viewer with a hosted dashboard when the audit volume grows.",
            "Add native OpenTelemetry SDK spans when deploying in a production service.",
            "Connect the webhook approval hook to your enterprise workflow system.",
            "Expand golden datasets for RAG grounding, citation validation, hallucination checks, and tool misuse detection.",
        ],
    )

    doc.add_section(WD_SECTION.CONTINUOUS)
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "PolicyAware AI Gateway User Guide"

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
